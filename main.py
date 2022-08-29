
import openpyxl as xl
from docx import Document
from docx.shared import Cm
from bs4 import BeautifulSoup

################################

XLSX_FILENAME = 'table.xlsx'
OUTPUT_FILENAME = 'output.docx'

CELL_WIDHT_PARAMS_NAME = Cm(4.5)
CELL_WIDHT_PARAMS = Cm(13)

IMG_FOLDER_PATH = 'img\\'

################################

def clear_markdown(in_str):
    soup = BeautifulSoup(in_str, "html.parser")
    return soup.get_text(separator=" ")

def make_table(rows, labels):
    table = output_docx.add_table(rows=rows, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    cells = table.columns[0].cells
    for row in range(0, rows):
        cells[row].text = labels[row]
        cells[row].paragraphs[0].runs[0].font.bold = True
        cells[row].width = CELL_WIDHT_PARAMS_NAME
    return table

def is_exception(label, value):
    if (label == 'CompanionFlag' or label == 'Fly' or label == 'Climb' or label == 'Burrow' or label == 'Swim' or label == 'Land' or label == 'OffenseNote' or label == 'BaseStatistics' or label == 'ExtractsPrepared' or label == 'AgeCategory' or label == 'DontUseRacialHD' or label == 'VariantParent') and value == '0':
        return True
    return False

# will not extract a label with a an empty or a NULL value
# row - number of current row
def extract_labels_and_data(sheet, row):
    labels = []
    data = []
    for col in range(1, sheet.max_column):
        value = str(sheet.cell(row, col).value)
        label = sheet.cell(row-1, col).value
        
        if label is not None:
            if value != 'None' and value != 'NULL' and value != '#ERROR!' or label == 'Treasure':
                if is_exception(label, value):
                    pass
                else:
                    labels.append(label)
                    if label == "FullText":
                        value = clear_markdown(value)
                    data.append(value)
        else:
            break
    # labels.append('Image')
    # data.append(IMG_FOLDER_PATH)
    return labels,data

output_docx = Document('template.docx')
input_workbook = xl.load_workbook(XLSX_FILENAME)
sheet = input_workbook['Bestiary']

# for row in range(2, 12, 2):
for row in range(2, sheet.max_row+2, 2):
    labels, data = extract_labels_and_data(sheet, row)
    output_table = make_table(len(labels), labels)
    cells = output_table.columns[1].cells
    for col in range (0, len(cells)):
        if col == len(cells)-1:
            cells[col].paragraphs[0].add_run()
            cells[col].paragraphs[0].runs[0].add_picture(data[col], width=Cm(3))
        else:
            cells[col].text = data[col]
        cells[col].width = CELL_WIDHT_PARAMS
        if col == 0:
            cells[col].paragraphs[0].style = output_docx.styles['Heading 1']
    output_docx.add_page_break()

#authors
output_docx.add_paragraph('Idea by CallMeJesusPLS', style='Authors')
output_docx.add_paragraph('Email: slavik2998@gmail.com', style='Authors')
output_docx.add_paragraph('Images by Some Juicy AI', style='Authors')
output_docx.add_paragraph('Typesetting by Hvunt', style='Authors')
output_docx.add_paragraph('Email: hvunt32@gmail.com', style='Authors')


output_docx.save(OUTPUT_FILENAME)
# convert('output.docx')